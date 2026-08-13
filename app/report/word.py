"""Word export (plan §10) — RTL Persian .docx from a done Report.

Uses python-docx; paragraphs are right-aligned, text set to Vazirmatn when
available on the client machine (falls back to Tahoma), font size 11pt.
"""
import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def report_to_docx(rep: dict[str, Any]) -> bytes:
    """rep: {"title", "intro", "sections": {key: {title, content}}}"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Vazirmatn"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Vazirmatn")

    h = doc.add_heading(rep.get("title", "گزارش چارت تولد"), level=0)
    _rtl(h)
    for run in h.runs:
        run.font.name = "Vazirmatn"

    intro = doc.add_paragraph(rep.get("intro", ""))
    _rtl(intro)

    for key, sec in (rep.get("sections") or {}).items():
        title = sec.get("title", key)
        content = sec.get("content", "")
        h2 = doc.add_heading(title, level=1)
        _rtl(h2)
        for run in h2.runs:
            run.font.name = "Vazirmatn"
        for para in str(content).split("\n\n"):
            p = doc.add_paragraph(para)
            _rtl(p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
